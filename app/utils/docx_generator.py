from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def format_markdown(text):
    # remove bold markdown
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return text


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph in a docx file."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")

    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)

    run.append(rpr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_paragraph_with_citations(doc, text, references):
    """Add a paragraph and convert [1][2] into real hyperlinks."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)

    # split text on citation patterns like [1]
    parts = re.split(r"(\[\d+\])", text)

    for part in parts:
        match = re.match(r"\[(\d+)\]", part)

        if match:
            num = int(match.group(1))
            ref = next((r for r in references if r.get("id") == num), None)

            if ref and ref.get("url"):
                add_hyperlink(paragraph, f"[{num}]", ref["url"])
            else:
                paragraph.add_run(f"[{num}]")
        else:
            if part.strip():
                paragraph.add_run(format_markdown(part))

    return paragraph


def process_text(doc, text, references):
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # handle markdown tables
        if "|" in line:
            table_data = []
            while i < len(lines) and "|" in lines[i]:
                row = [cell.strip() for cell in lines[i].split("|") if cell.strip()]
                if row and not all(set(cell) <= {"-"} for cell in row):
                    table_data.append(row)
                i += 1

            if table_data:
                rows = len(table_data)
                cols = len(table_data[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = "Table Grid"

                for r in range(rows):
                    for c in range(cols):
                        if c < len(table_data[r]):
                            table.rows[r].cells[c].text = format_markdown(table_data[r][c])

            continue

        # subheadings
        elif line.startswith("### "):
            doc.add_heading(format_markdown(line.replace("### ", "")), level=3)

        elif line.startswith("## "):
            doc.add_heading(format_markdown(line.replace("## ", "")), level=2)

        # bullet points
        elif line.strip().startswith("- "):
            doc.add_paragraph(format_markdown(line.strip()[2:]), style="List Bullet")

        # normal paragraph with citation links
        elif line.strip():
            add_paragraph_with_citations(doc, line.strip(), references)

        i += 1


def generate_docx(topic: str, report: dict, filename="research_paper.docx"):
    doc = Document()

    
    doc.add_heading(topic, 0)
    doc.add_paragraph("AI Research Assistant")

    references = report.get("references", [])

    section_number = 1

    for section, text in report.items():
        if section == "references":
            continue

        if not isinstance(text, str):
            continue

        doc.add_heading(f"{section_number}. {section}", level=1)
        section_number += 1

        process_text(doc, text, references)

    
    if references:
        doc.add_heading("References", level=1)
        for ref in references:
            para = doc.add_paragraph()
            para.add_run(f"[{ref['id']}] {ref.get('title', 'Untitled')} — ")

            if ref.get("url"):
                add_hyperlink(para, ref["url"], ref["url"])

    doc.save(filename)
    return filename
